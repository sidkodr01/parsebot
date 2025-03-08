import os
import requests
import fitz  # For PDF processing
from bs4 import BeautifulSoup  # For website scraping
from docx import Document as DocxDocument  # For Word document processing
from langchain.text_splitter import CharacterTextSplitter
from langchain_google_genai import ChatGoogleGenerativeAI, GoogleGenerativeAIEmbeddings
from langchain_community.vectorstores import FAISS
from langchain.chains import RetrievalQA
from langchain.prompts import PromptTemplate
from langchain.memory import ConversationBufferMemory
from langchain.docstore.document import Document
import tempfile
from dotenv import load_dotenv

load_dotenv()

# Configuration variables
CHUNK_SIZE = 300
CHUNK_OVERLAP = 50
MAX_TOKENS = 15000
MODEL_NAME = "gemini-2.0-flash"  # Google's model name
TEMPERATURE = 0.4

# Set up Google API key
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")
if not GOOGLE_API_KEY:
    raise ValueError("Please set the GOOGLE_API_KEY in the .env file.")


# ================== PDF Processing ==================
def read_pdf(file_path):
    """Reads a PDF file and extracts text content."""
    try:
        doc = fitz.open(file_path)
        return [page.get_text().strip() for page in doc if page.get_text().strip()]
    except Exception as e:
        print(f"Error reading the PDF file: {e}")
        return []

def process_pdf(file_path):
    """Processes a PDF file and splits it into chunks."""
    content = read_pdf(file_path)
    if not content:
        raise ValueError("No content could be read from the PDF file.")
    documents = [Document(page_content=text, metadata={"source": file_path}) for text in content]
    text_splitter = CharacterTextSplitter(chunk_size=CHUNK_SIZE, chunk_overlap=CHUNK_OVERLAP)
    return text_splitter.split_documents(documents)

# ================== Website Processing ==================
def scrape_website(url):
    """Scrapes content from a website."""
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
    }
    try:
        response = requests.get(url, headers=headers)
        response.raise_for_status()
        soup = BeautifulSoup(response.text, 'html.parser')
        
        # Remove script and style elements
        for script in soup(["script", "style"]):
            script.decompose()
        
        # Get text from various elements
        content = []
        for elem in soup.find_all(['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'li', 'span', 'div']):
            if elem.text.strip():
                content.append(elem.text.strip())
        
        # If no content found, try to get all text from body
        if not content:
            body = soup.find('body')
            if body:
                content = [body.get_text(separator='\n', strip=True)]
        
        if not content:
            print("Warning: No content found. The website might have unusual structure or require JavaScript.")
            return []
        
        return content
    except requests.RequestException as e:
        print(f"Error scraping the website: {e}")
        return []

def process_website(url):
    """Processes website content and splits it into chunks."""
    content = scrape_website(url)
    if not content:
        raise ValueError("No content could be fetched from the website.")
    documents = [Document(page_content=text, metadata={"source": url}) for text in content]
    text_splitter = CharacterTextSplitter(chunk_size=CHUNK_SIZE, chunk_overlap=CHUNK_OVERLAP)
    return text_splitter.split_documents(documents)

# ================== Word Document Processing ==================
def read_word_document(file_path):
    """Reads a Word document and extracts text content."""
    try:
        doc = DocxDocument(file_path)
        return [paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip()]
    except Exception as e:
        print(f"Error reading the Word document: {e}")
        return []

def process_word_document(file_path):
    """Processes a Word document and splits it into chunks."""
    content = read_word_document(file_path)
    if not content:
        raise ValueError("No content could be read from the Word document.")
    documents = [Document(page_content=text, metadata={"source": file_path}) for text in content]
    text_splitter = CharacterTextSplitter(chunk_size=CHUNK_SIZE, chunk_overlap=CHUNK_OVERLAP)
    return text_splitter.split_documents(documents)

# ================== RAG Pipeline ==================
def initialize_rag_pipeline(texts):
    """Initializes the RAG pipeline with the given texts."""
    embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")
    vectorstore = FAISS.from_documents(texts, embeddings)
    
    # Set up the retrieval-based QA system with a prompt template
    template = """Context: {context}

Question: {question}

Answer the question concisely based only on the given context. If the context doesn't contain relevant information, say "I don't have enough information to answer that question."

But, if the question is generic, then go ahead and answer the question, example what is a electric vehicle?

Answer the question as human like as possible.


"""
    PROMPT = PromptTemplate(template=template, input_variables=["context", "question"])
    memory = ConversationBufferMemory(memory_key="chat_history", return_messages=True)
    
    qa = RetrievalQA.from_chain_type(
        llm=ChatGoogleGenerativeAI(model=MODEL_NAME, temperature=TEMPERATURE, max_tokens=MAX_TOKENS),
        chain_type="stuff",
        retriever=vectorstore.as_retriever(),
        memory=memory,
        chain_type_kwargs={"prompt": PROMPT}
    )
    return qa, vectorstore

def rag_pipeline(query, qa_chain, vectorstore):
    """Runs the RAG pipeline to answer a query."""
    relevant_docs = vectorstore.similarity_search_with_score(query, k=3)
    if not relevant_docs:
        return "I don't have enough information to answer that question."
    context = "\n\n".join([doc.page_content for doc, _ in relevant_docs])
    response = qa_chain.invoke({"query": query})
    return response['result']

# ================== Main Program ==================
if __name__ == "__main__":
    print("Welcome to the Unified Chatbot!")
    print("This chatbot can process PDFs, websites, and Word documents.")
    
    while True:
        print("\nChoose an option:")
        print("1. Process a PDF file")
        print("2. Process a website")
        print("3. Process a Word document")
        print("4. Quit")
        choice = input("Enter your choice (1/2/3/4): ").strip()
        
        if choice == "4":
            print("Exiting the program. Goodbye!")
            break
        
        if choice not in ["1", "2", "3"]:
            print("Invalid choice. Please try again.")
            continue
        
        try:
            if choice == "1":
                file_path = input("Enter the path to the PDF file: ").strip()
                if not os.path.exists(file_path):
                    print("File not found. Please try again.")
                    continue
                print("Processing PDF file...")
                texts = process_pdf(file_path)
            
            elif choice == "2":
                url = input("Enter the URL of the website: ").strip()
                print("Processing website content...")
                texts = process_website(url)
            
            elif choice == "3":
                file_path = input("Enter the path to the Word document: ").strip()
                if not os.path.exists(file_path):
                    print("File not found. Please try again.")
                    continue
                print("Processing Word document...")
                texts = process_word_document(file_path)
            
            if not texts:
                print("No content found. Please try a different file or URL.")
                continue
            
            print("Initializing RAG pipeline...")
            qa, vectorstore = initialize_rag_pipeline(texts)
            
            print("\nRAG Pipeline initialized. You can now enter your queries.")
            print("Enter 'new' to process a new file/website or 'quit' to exit.")
            
            while True:
                user_query = input("\nEnter your query: ").strip()
                if user_query.lower() == 'quit':
                    print("Exiting the program. Goodbye!")
                    exit()
                elif user_query.lower() == 'new':
                    break
                
                result = rag_pipeline(user_query, qa, vectorstore)
                print(f"Response: {result}")
        
        except Exception as e:
            print(f"An error occurred: {e}")
            print("Please try again.")